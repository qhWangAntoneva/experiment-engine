"""Chinese Word (.docx) report generation for QCA analysis results.

Generates a formatted .docx report using python-docx. The library is
lazily installed via micropip at runtime in the Pyodide browser
environment -- it is NOT a project-level dependency.
"""

from __future__ import annotations

import io


class QCADocxReporter:
    """Generate Chinese Word (.docx) reports from QCA analysis results.

    Usage::

        reporter = QCADocxReporter()
        docx_bytes = reporter.generate(result, robustness=robustness_report)
        # docx_bytes can be saved to a file or returned as a browser download.

    Sections:
    1. Title page (title, date, case count, condition count, outcome name)
    2. Executive summary (Chinese NL interpretation)
    3. Truth table (formatted table)
    4. Solutions (all 3 types with formulas, consistency, coverage)
    5. Necessity analysis (table)
    6. Robustness (if provided)
    7. Charts (embedded PNG images, if provided)
    """

    # ── Constants ─────────────────────────────────────────────────────────

    HEADING_FONT = "SimHei"
    BODY_FONT = "SimSun"
    TABLE_STYLE = "Table Grid"

    @staticmethod
    def _set_run_font(run, font_name: str, size_pt: float = 10.5, bold: bool = False):
        """Set both Western and East-Asian font on a run."""
        run.font.name = font_name
        run.font.size = int(size_pt * 12700)  # EMU: 1 pt = 12700 EMU
        run.bold = bold
        # Set east-Asia font via XML
        rpr = run._element.get_or_add_rPr()  # noqa: SLF001
        from docx.oxml.ns import (
            qn,
        )

        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            from lxml import etree

            rfonts = etree.SubElement(rpr, qn("w:rFonts"))
        rfonts.set(qn("w:eastAsia"), font_name)

    @staticmethod
    def _add_heading(doc, text: str, level: int = 1):
        """Add a heading with SimHei font."""
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            QCADocxReporter._set_run_font(run, QCADocxReporter.HEADING_FONT, bold=True)

    @staticmethod
    def _add_paragraph(doc, text: str, bold: bool = False):
        """Add a body paragraph with SimSun font."""
        p = doc.add_paragraph()
        run = p.add_run(text)
        QCADocxReporter._set_run_font(run, QCADocxReporter.BODY_FONT, bold=bold)

    @staticmethod
    def _add_table_row(table, cells_text: list[str], bold: bool = False):
        """Add a row to a table with proper fonts."""

        row = table.add_row()
        for i, text in enumerate(cells_text):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            QCADocxReporter._set_run_font(run, QCADocxReporter.BODY_FONT, bold=bold)

    # ── Executive Summary helpers (Chinese NL) ────────────────────────────

    @staticmethod
    def _cn_numeral(n: int) -> str:
        cn = ["零", "一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
        if 0 <= n < len(cn):
            return cn[n]
        return str(n)

    @staticmethod
    def _solution_type_cn(st: str) -> str:
        return {
            "complex": "复杂解",
            "parsimonious": "精简解",
            "intermediate": "中间解",
        }.get(st, st)

    @staticmethod
    def _consistency_quality(c: float) -> str:
        if c >= 0.95:
            return "非常高"
        if c >= 0.9:
            return "很高"
        if c >= 0.8:
            return "较高"
        if c >= 0.75:
            return "可接受"
        return "偏低"

    @staticmethod
    def _coverage_quality(c: float) -> str:
        if c >= 0.8:
            return "很高"
        if c >= 0.6:
            return "较高"
        if c >= 0.4:
            return "中等"
        return "较低"

    @staticmethod
    def _generate_summary(result) -> str:
        """Generate Chinese natural-language executive summary."""
        sols = result.solutions
        primary = sols.intermediate or sols.complex or sols.parsimonious
        if primary is None or not primary.terms:
            return "未找到有效的解。"

        outcome = result.fuzzy_data.outcome_name if result.fuzzy_data else "结果"
        n_terms = len(primary.terms)

        lines = []
        lines.append(f"=     =   {outcome} =  QCA     =")

        if n_terms == 1:
            lines.append(f'          "{outcome}"                 ')
        else:
            lines.append(
                f'          "{outcome}"     {QCADocxReporter._cn_numeral(n_terms)}         '
            )

        path_labels = ["一", "二", "三", "四", "五", "六"]
        for i, term in enumerate(primary.terms):
            idx = path_labels[i] if i < len(path_labels) else str(i + 1)
            label = term.label or " * ".join(term.term) if term.term else "未知条件组合"
            lines.append(f"     {idx}    {label}")

        sol_type_cn = QCADocxReporter._solution_type_cn(primary.solution_type)
        lines.append("")
        lines.append(f"       {sol_type_cn} =")

        c = primary.solution_consistency
        q = QCADocxReporter._consistency_quality(c)
        lines.append(f"=          {c:.3f} = {q} =")

        cov = primary.solution_coverage
        qc = QCADocxReporter._coverage_quality(cov)
        lines.append(f"=         {cov:.3f} = {qc} =")

        return "\n".join(lines)

    # ── Main generation ───────────────────────────────────────────────────

    def generate(
        self,
        result,
        robustness=None,
        title: str = "QCA    ",
        charts: dict[str, bytes] | None = None,
    ) -> bytes:
        """Generate a complete QCA .docx report.

        Args:
            result: QCAAnalysisResult instance.
            robustness: Optional RobustnessReport.
            title: Document title.
            charts: Optional dict mapping chart names to PNG bytes (not yet
                    embedded -- placeholder for future integration).

        Returns:
            Report as bytes (ready for download).
        """
        from datetime import datetime

        from docx import Document

        _ = charts  # reserved for future chart embedding

        doc = Document()

        # Use a compact default style
        style = doc.styles["Normal"]
        style.font.name = self.BODY_FONT

        # ── 1. Title page ─────────────────────────────────────────────────
        n_cases = result.fuzzy_data.n_cases if result.fuzzy_data else 0
        n_conds = result.fuzzy_data.n_conditions if result.fuzzy_data else 0
        outcome_name = result.fuzzy_data.outcome_name if result.fuzzy_data else "N/A"

        doc.add_paragraph()  # spacing
        self._add_heading(doc, title, level=0)
        self._add_paragraph(doc, f"    {datetime.now().strftime('%Y-%m-%d')}")
        self._add_paragraph(doc, "")
        self._add_paragraph(doc, f"    {n_cases}")
        self._add_paragraph(doc, f"    {n_conds}")
        self._add_paragraph(doc, f"    {outcome_name}")
        doc.add_page_break()

        # ── 2. Executive summary ──────────────────────────────────────────
        self._add_heading(doc, "    ", level=1)
        summary_text = self._generate_summary(result)
        for line in summary_text.split("\n"):
            self._add_paragraph(doc, line)

        # ── 3. Truth table ────────────────────────────────────────────────
        if result.truth_table:
            doc.add_page_break()
            self._add_heading(doc, "    ", level=1)
            tt = result.truth_table
            self._add_paragraph(
                doc,
                f"        {tt.consistency_threshold}          {tt.frequency_threshold}",
            )
            self._add_paragraph(doc, "")

            rows = tt.included_rows
            if rows:
                table = doc.add_table(rows=1, cols=4)
                table.style = self.TABLE_STYLE
                # Header
                hdr_cells = table.rows[0].cells
                for i, label in enumerate(["  ", "  ", "    ", "  "]):
                    hdr_cells[i].text = ""
                    p = hdr_cells[i].paragraphs[0]
                    run = p.add_run(label)
                    self._set_run_font(run, self.HEADING_FONT, bold=True)

                for r in rows:
                    self._add_table_row(
                        table,
                        [
                            r.config_label,
                            f"{r.frequency:.1f}",
                            f"{r.raw_consistency:.3f}",
                            "1" if r.outcome_value else "0",
                        ],
                    )

        # ── 4. Solutions ──────────────────────────────────────────────────
        if result.solutions:
            doc.add_page_break()
            self._add_heading(doc, "QCA  ", level=1)

            for sol_type in ("complex", "parsimonious", "intermediate"):
                sol = getattr(result.solutions, sol_type, None)
                if sol is None or not sol.terms:
                    continue

                cn_label = self._solution_type_cn(sol_type)
                self._add_heading(doc, cn_label, level=2)
                self._add_paragraph(doc, f"    {sol.formula}")
                self._add_paragraph(doc, f"        {sol.solution_consistency:.3f}")
                self._add_paragraph(doc, f"        {sol.solution_coverage:.3f}")

                if sol.terms:
                    term_table = doc.add_table(rows=1, cols=4)
                    term_table.style = self.TABLE_STYLE
                    hdr = term_table.rows[0].cells
                    for j, lbl in enumerate(["    ", "    ", "    ", "    "]):
                        hdr[j].text = ""
                        p = hdr[j].paragraphs[0]
                        run = p.add_run(lbl)
                        self._set_run_font(run, self.HEADING_FONT, bold=True)
                    for term in sol.terms:
                        self._add_table_row(
                            term_table,
                            [
                                term.label or " * ".join(term.term)
                                if term.term
                                else "-",
                                f"{term.consistency:.3f}",
                                f"{term.raw_coverage:.3f}",
                                f"{term.unique_coverage:.3f}",
                            ],
                        )
                self._add_paragraph(doc, "")

        # ── 5. Necessity analysis ─────────────────────────────────────────
        if result.necessity:
            doc.add_page_break()
            self._add_heading(doc, "    ", level=1)
            nec = result.necessity
            self._add_paragraph(doc, f"    {nec.threshold}")
            self._add_paragraph(doc, "")

            nec_table = doc.add_table(rows=1, cols=4)
            nec_table.style = self.TABLE_STYLE
            hdr = nec_table.rows[0].cells
            for j, lbl in enumerate(["  ", "    ", "  ", "    "]):
                hdr[j].text = ""
                p = hdr[j].paragraphs[0]
                run = p.add_run(lbl)
                self._set_run_font(run, self.HEADING_FONT, bold=True)
            for c in nec.conditions:
                self._add_table_row(
                    nec_table,
                    [
                        c.condition_name,
                        f"{c.consistency:.3f}",
                        f"{c.coverage:.3f}",
                        "是" if c.is_necessary else "否",
                    ],
                )

        # ── 6. Robustness ─────────────────────────────────────────────────
        if robustness is not None:
            doc.add_page_break()
            self._add_heading(doc, "    ", level=1)
            self._add_paragraph(
                doc,
                f"      {robustness.overall_robustness:.2f}",
            )
            self._add_paragraph(doc, robustness.summary)
            self._add_paragraph(doc, "")

            rob_table = doc.add_table(rows=1, cols=4)
            rob_table.style = self.TABLE_STYLE
            hdr = rob_table.rows[0].cells
            for j, lbl in enumerate(["    ", "    ", "    ", "  "]):
                hdr[j].text = ""
                p = hdr[j].paragraphs[0]
                run = p.add_run(lbl)
                self._set_run_font(run, self.HEADING_FONT, bold=True)
            for t in robustness.tests:
                if t.solution_stability:
                    stability_str = (
                        f"{t.solution_stability[0]:.3f}-{t.solution_stability[-1]:.3f}"
                    )
                else:
                    stability_str = "N/A"
                self._add_table_row(
                    rob_table,
                    [
                        t.test_name,
                        t.parameter_varied,
                        stability_str,
                        "通过" if t.passed else "失败",
                    ],
                )

        # ── Write to BytesIO ──────────────────────────────────────────────
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()
